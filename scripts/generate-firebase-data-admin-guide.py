#!/usr/bin/env python3
"""Generate the Moonly Firebase web admin guide as a DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "web-admin" / "moonly-firebase-data-admin-guide.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 35, 40)
MUTED = RGBColor(104, 112, 118)
HEADER_FILL = "E8EEF5"
SOFT_FILL = "F4F6F9"
LINE = "D9D5CC"


PAGES = [
    {
        "title": "后台界面和通用操作",
        "lead": "后台用于让管理员在浏览器中读取、筛选、新增、修改和删除 Firestore 数据。界面采用三栏结构：左侧选择数据源和查询条件，中间展示文档列表，右侧编辑单个文档 JSON。",
        "table": [
            ("读取", "选择集合、设置排序和数量后点击刷新；文档 ID 可以精确定位单条记录"),
            ("新增", "点击新增，填写文档 ID 或留空自动生成，再写入 JSON"),
            ("修改", "点选文档后在右侧 JSON 中修改字段，默认以 merge 方式合并字段"),
            ("删除", "点选文档后删除；勾选连带子集合时使用 recursiveDelete"),
        ],
        "bullets": [
            "所有写操作都会直接影响线上 Firestore，生产环境建议先在测试集合或备份数据上演练。",
            "where 查询支持 ==、!=、<、<=、>、>= 和 array-contains，值会尝试按 JSON 解析。",
            "导出 JSON 会导出当前列表中已经加载的文档，不代表整个集合的全量备份。",
            "复制路径功能用于把精确文档路径贴到日志、工单或 Firebase Console 中交叉核验。",
        ],
    },
    {
        "title": "管理员登录和权限模型",
        "lead": "没有默认管理员账号。管理员账号必须先在 Firebase Authentication 中创建，并在 custom claim 或 Firestore 用户文档里授予管理员标记。",
        "table": [
            ("可用凭据", "Firebase Auth 邮箱密码账号"),
            ("Claim 权限", "admin=true 或 role=\"admin\""),
            ("Firestore 权限", "users/{uid}.isAdmin=true 或 users/{uid}.role=\"admin\""),
            ("失败表现", "未登录返回 401；不是管理员返回 403；路径或参数错误返回 400"),
        ],
        "bullets": [
            "密码不能从项目代码中找回；忘记密码时应在 Firebase Console 里重置。",
            "如果只在 Firestore 添加 role 字段，要确保该 UID 对应的用户文档已经存在。",
            "管理员离职或外包结束后，应同时禁用 Auth 用户和移除 Firestore 管理员标记。",
        ],
    },
    {
        "title": "根集合总览",
        "lead": "系统内置了当前项目代码和规则中出现的根集合，同时会通过 Admin SDK 发现其他根集合。发现的新集合会以自定义路径方式读取。",
        "table": [
            ("users / public_profiles", "账号资料和公开搜索资料，可读取和修改"),
            ("feedback / app_config / quick_reading", "反馈处理、公开配置、快速占卜配置，可读写"),
            ("daily_oracle_summaries / relationship_divinations", "好友动态摘要和双人关系占卜，可读写"),
            ("iap_receipts / usage_limits / payment_events / analytics_events", "付费、用量、支付事件和分析事件，建议谨慎修改"),
        ],
        "bullets": [
            "根集合不需要输入用户 UID，直接选择集合即可读取。",
            "如果集合默认排序字段不存在，后端会回退到不排序查询以避免后台完全不可用。",
            "修改分析、支付、收据类数据前，建议先导出当前列表作为人工回滚参考。",
        ],
    },
    {
        "title": "users 用户资料",
        "lead": "users/{uid} 是用户核心资料文档，也是管理员权限判断、会员状态、资料同步和账号删除的中心。",
        "table": [
            ("路径", "users/{firebaseUid}"),
            ("可读取", "displayName、email、photoUrl、birthday、city、timezone、membershipStatus、proExpiresAt"),
            ("可修改", "资料字段、会员字段、isAdmin、role、profileUpdatedAt 等"),
            ("风险", "错误修改 UID 对应文档可能影响登录后的资料、会员判断和搜索索引"),
        ],
        "bullets": [
            "给管理员授权时，推荐只添加 isAdmin=true 或 role=\"admin\"，不要改动用户登录字段。",
            "会员字段 membershipStatus 和 proExpiresAt 会影响 membershipStatus Function 返回结果。",
            "删除用户资料前，确认是否还需要删除 public_profiles 和用户子集合中的历史数据。",
        ],
    },
    {
        "title": "public_profiles 公开资料",
        "lead": "public_profiles/{uid} 用于好友搜索和推荐，保留公开展示所需的轻量资料。它通常应与 users/{uid} 的姓名、头像和搜索字段保持一致。",
        "table": [
            ("路径", "public_profiles/{uid}"),
            ("可读取", "displayName、displayNameLower、email、photoUrl、searchKeywords"),
            ("可修改", "昵称、头像、搜索关键字和公开简介字段"),
            ("风险", "displayNameLower 或 searchKeywords 不正确会导致搜索不到用户"),
        ],
        "bullets": [
            "批量修复搜索问题时，优先检查 displayNameLower 是否是小写/规范化结果。",
            "公开资料不应写入生日、支付、私聊、记忆等敏感或私有字段。",
            "如果用户改名，建议同时检查 users/{uid} 与 public_profiles/{uid} 的同步状态。",
        ],
    },
    {
        "title": "feedback 反馈镜像",
        "lead": "反馈会写入 users/{uid}/feedback/{feedbackId} 和顶层 feedback/{feedbackId}。通用后台能直接管理镜像，原反馈后台仍可专门处理状态流。",
        "table": [
            ("路径", "feedback/{feedbackId}"),
            ("可读取", "content、status、category、tag、uid、email、platform、appVersion、deviceModel"),
            ("可修改", "status、adminNote、handledBy、分类标签和备注"),
            ("风险", "只改顶层镜像不会自动同步用户子集合，专门反馈后台会做双写同步"),
        ],
        "bullets": [
            "处理状态建议使用 new、triaged、in_progress、resolved、closed。",
            "如果要让 App 内用户也看到处理结果，应使用专门反馈后台或同时修改用户子集合。",
            "导出反馈 JSON 可以作为客服日报、问题分类或版本回归分析的原始材料。",
        ],
    },
    {
        "title": "app_config 和 quick_reading",
        "lead": "app_config/public 控制公开链接和 IAP 商品配置，quick_reading/{oracleId} 控制快速占卜内容。两者都直接影响客户端展示。",
        "table": [
            ("app_config/public", "socialLinks、iapProducts、priceLabel、productId"),
            ("quick_reading/{oracleId}", "按 tarot、astrology、sage 等 oracleId 存放快速占卜配置"),
            ("可修改", "社媒链接、商品 ID、价格标签、快速占卜文案或配置"),
            ("风险", "商品 ID 改错会影响 IAP；快速占卜配置改错会影响 App 启动后的内容读取"),
        ],
        "bullets": [
            "公开配置修改后，建议用 publicConfig Function 或 App 端重新打开页面验证。",
            "IAP 商品 ID 必须与 App Store / Google Play / Unity IAP 配置一致。",
            "quick_reading 是只读客户端集合，管理员通过后台修改时要保留原有文档 ID。",
        ],
    },
    {
        "title": "daily_oracle_summaries 每日神谕摘要",
        "lead": "daily_oracle_summaries/{uid}_{date} 是朋友动态使用的摘要，不包含完整私密解读。它由用户每日神谕记录同步生成。",
        "table": [
            ("路径", "daily_oracle_summaries/{uid}_{yyyy-MM-dd}"),
            ("可读取", "ownerUid、date、cardId、cardName、orientation、title、oracle、visibility"),
            ("可修改", "syncEnabled、visibility、summaryOnly、oracle 文案和 updatedAt"),
            ("风险", "summaryOnly 必须保持 true；visibility 错误会影响好友可见性"),
        ],
        "bullets": [
            "ownerUid 与文档 ID 前缀应一致，便于排查某个用户的公开摘要。",
            "如果用户关闭同步，应将 syncEnabled=false 且 visibility=\"only_me\"。",
            "摘要只应该放短文案，不应把 detail、dos、donts 等完整内容搬到顶层摘要。",
        ],
    },
    {
        "title": "relationship_divinations 双人关系占卜",
        "lead": "relationship_divinations/{readingId} 保存双人关系占卜邀请、双方状态和抽牌结果。它与用户好友关系和个人历史有交叉影响。",
        "table": [
            ("路径", "relationship_divinations/{readingId}"),
            ("可读取", "initiatorUid、receiverUid、question、status、cards、createdAt、completedAt"),
            ("可修改", "status、initiatorRevealed、receiverJoined、receiverRevealed、completedAt"),
            ("风险", "cards 不应在进行中随意改动；状态字段不一致会让双方看到不同进度"),
        ],
        "bullets": [
            "常见状态包括 invited、initiator_revealed、receiver_joined、completed、cancelled。",
            "如果双方都已翻牌，status 应为 completed，并写入 completedAt。",
            "管理员介入取消时，推荐只改 status=cancelled，不删除记录，保留审计线索。",
        ],
    },
    {
        "title": "IAP、会员和支付事件",
        "lead": "iap_receipts、payment_events 与 users/{uid} 的会员字段共同决定付费结果。它们可以读取和修正，但生产环境要谨慎。",
        "table": [
            ("iap_receipts/{receiptId}", "productId、store、status、valid、proExpiresAt、receiptHash"),
            ("payment_events/{eventId}", "uid、provider、transactionId、membershipStatus、payload"),
            ("users/{uid}", "membershipStatus、proExpiresAt、membershipProvider、membershipTransactionId"),
            ("风险", "手工把用户改成 pro 可能绕过真实收据校验；过期时间格式必须正确"),
        ],
        "bullets": [
            "真实收据校验仍应优先使用 submitIapReceipt Function，而不是直接改 Firestore。",
            "修复会员时，要同步检查 users/{uid} 和最近一条 iap_receipts 状态。",
            "payload 可能包含支付平台返回信息，导出和分享时要注意隐私。",
        ],
    },
    {
        "title": "usage_limits 和 analytics_events",
        "lead": "usage_limits 控制 AI/TTS 每日用量，analytics_events 记录客户端事件。后台可以读取这些集合，也可以在必要时修正异常计数。",
        "table": [
            ("usage_limits/{uid_action_day}", "uid、action、day、count、limit、updatedAt"),
            ("analytics_events/{eventId}", "由客户端创建的分析事件，默认客户端不能读取"),
            ("可修改", "用量 count、limit、refundedCount；分析事件一般只读或删除异常数据"),
            ("风险", "随意降低 count 会让用户重新获得额度；分析事件不可作为强一致业务状态"),
        ],
        "bullets": [
            "AI 动作通常是 aiChat，TTS 动作通常是 tts。",
            "用量修正适合处理接口失败但未退还额度的客服场景。",
            "分析事件建议只用于趋势和排障，不用于判定会员或交易。",
        ],
    },
    {
        "title": "用户子集合：daily_oracles",
        "lead": "users/{uid}/daily_oracles/{date} 保存用户自己的每日神谕完整记录，比顶层摘要包含更多私密内容。",
        "table": [
            ("路径", "users/{uid}/daily_oracles/{yyyy-MM-dd}"),
            ("可读取", "cardId、cardName、orientation、title、oracle、detail、dos、donts、microAction"),
            ("可修改", "完整牌面、文案、locale、oracleId、syncEnabled、visibility"),
            ("风险", "如果同步到好友动态，还要检查 daily_oracle_summaries 是否一致"),
        ],
        "bullets": [
            "date 字段和文档 ID 应保持同一天，避免近期记录排序或查询异常。",
            "dos、donts 是数组，JSON 编辑时应保持为字符串数组。",
            "用户投诉某日内容错误时，先查用户子集合，再看摘要是否已发布。",
        ],
    },
    {
        "title": "用户子集合：divination_records",
        "lead": "users/{uid}/divination_records/{readingId} 是占卜历史记录，包含问题、牌阵、抽到的牌、摘要和完整解读字段。",
        "table": [
            ("路径", "users/{uid}/divination_records/{readingId}"),
            ("可读取", "question、scene、spreadKind、lockedCards、shortVerdict、judgeContent、adviceContent"),
            ("可修改", "问题、场景、解读文案、topics、oracleId、createdAt、updatedAt"),
            ("风险", "lockedCards 是对象数组；结构改坏会导致历史详情页无法正确显示"),
        ],
        "bullets": [
            "lockedCards 每项通常包含 positionKey、position、cardId、cardName、orientation。",
            "createdAt 如果写成时间戳包装对象，后台会还原成 Firestore Timestamp。",
            "删除历史记录会影响用户历史页，但不会删除对话中的上下文内容。",
        ],
    },
    {
        "title": "用户子集合：dialog_sessions",
        "lead": "users/{uid}/dialog_sessions/default 保存默认对话会话，包含 UI 消息、API 消息、活动占卜上下文和附件。",
        "table": [
            ("路径", "users/{uid}/dialog_sessions/{sessionId}，当前常用 default"),
            ("可读取", "messages、apiMessages、activeReadingId、activeReadingState、activeContextAttachments"),
            ("可修改", "对话快照、当前活动状态、附件摘要和 savedAtLocal"),
            ("风险", "消息结构复杂；错误修改 roleType、messageType 或 cards 数据会影响聊天恢复"),
        ],
        "bullets": [
            "一般排障建议先导出 JSON，再只改 activeReadingId 等小范围字段。",
            "不建议在生产环境手工编辑长 messages 数组，除非明确知道字段结构。",
            "用户要求清空聊天时，可删除 default 文档或把 messages/apiMessages 置为空数组。",
        ],
    },
    {
        "title": "用户子集合：memories 和 tomorrow_hooks",
        "lead": "memories/runtime 保存用户记忆或偏好快照，tomorrow_hooks 保存明日提醒、后续关怀或待触发内容。",
        "table": [
            ("users/{uid}/memories/runtime", "运行时记忆、隐私设置相关内容和更新时间"),
            ("users/{uid}/tomorrow_hooks/{hookId}", "提醒标题、内容、触发日期、来源和状态"),
            ("可修改", "记忆内容、启用状态、hook 状态、触发时间和备注字段"),
            ("风险", "记忆类数据较敏感；导出或复制时应最小化暴露"),
        ],
        "bullets": [
            "memory_privacy 设置和 memories 内容要一起看，避免用户关闭记忆后仍保留旧内容。",
            "tomorrow_hooks 用于后续体验，不应写入支付、身份凭据或管理员备注。",
            "删除记忆前，建议确认用户是否只是想关闭后续使用，而不是完全删除历史。",
        ],
    },
    {
        "title": "用户子集合：好友关系",
        "lead": "friends、friend_requests、blocked_users 共同描述社交关系。管理员可以读取和修正单边异常，但要注意双方镜像关系。",
        "table": [
            ("friends/{friendUid}", "好友 UID、显示名、头像、status、updatedAt"),
            ("friend_requests/{requesterUid}", "请求方资料、状态、时间和附加信息"),
            ("blocked_users/{blockedUid}", "拉黑 UID、blockedAt、reason 等"),
            ("风险", "只改一方 friends 可能造成 A 看见 B、B 看不见 A 的不一致"),
        ],
        "bullets": [
            "恢复好友关系时，需要检查双方 users/{uid}/friends/{otherUid}。",
            "处理好友请求卡住时，先查 receiver 的 friend_requests，再查 requester 的 friends。",
            "拉黑操作通常应同时删除双方 friends 和 friend_requests。",
        ],
    },
    {
        "title": "用户子集合：virtual_friends、settings、user_feedback",
        "lead": "virtual_friends 是用户自建好友档案；settings 存放通知、同步和记忆隐私设置；user_feedback 是用户侧反馈记录。",
        "table": [
            ("virtual_friends/{id}", "姓名、关系、背景、备注、头像和更新时间"),
            ("settings/notifications", "通知开关、时间、渠道偏好"),
            ("settings/daily_divination_sync / memory_privacy", "每日同步可见性、记忆隐私控制"),
            ("feedback/{feedbackId}", "用户自己的反馈副本，最好与顶层 feedback 镜像保持一致"),
        ],
        "bullets": [
            "自建好友属于用户私有内容，后台排障时只读取必要字段。",
            "同步设置会影响 daily_oracle_summaries 是否继续发布。",
            "如果顶层反馈状态已更新但用户端没显示，检查用户子集合反馈副本。",
        ],
    },
    {
        "title": "JSON 编辑格式和特殊 Firestore 类型",
        "lead": "后台右侧编辑器使用 JSON。为了保留 Firestore 特殊类型，系统支持带 __type 的包装对象，并在保存时转换回 Admin SDK 类型。",
        "table": [
            ("Timestamp", "{\"__type\":\"timestamp\",\"value\":\"2026-06-23T10:00:00Z\"}"),
            ("ServerTimestamp", "{\"__type\":\"serverTimestamp\"}"),
            ("DeleteField", "{\"__type\":\"deleteField\"}"),
            ("GeoPoint / Reference", "{\"__type\":\"geoPoint\",\"latitude\":31.2,\"longitude\":121.5}；{\"__type\":\"reference\",\"path\":\"users/uid\"}"),
        ],
        "bullets": [
            "普通字符串、数字、布尔值、数组和对象会按 JSON 原样保存。",
            "合并字段开启时，未出现的字段不会被删除；要删除字段请使用 deleteField 包装对象。",
            "关闭合并字段会覆盖整个文档，适合重建配置，不适合修补用户资料。",
        ],
    },
    {
        "title": "上线、部署和验收清单",
        "lead": "功能完成后，需要同时部署 Functions 和 Hosting。部署前应完成语法检查、页面加载检查、管理员权限检查和至少一次测试集合 CRUD 演练。",
        "table": [
            ("部署命令", "firebase deploy --only functions,hosting --project fari-app-b2fd2"),
            ("本地验证", "npm --prefix functions run lint；打开 /data-admin.html 无控制台错误"),
            ("权限验证", "普通登录用户应返回 403；管理员用户应能列集合"),
            ("数据验证", "测试集合新增、读取、修改、删除均成功，生产集合先导出再修改"),
        ],
        "bullets": [
            "functions/ 目录被全局 gitignore 忽略，提交时需要 git add -f functions/index.js。",
            "20 页手册放在 web-admin 下，部署后可通过后台页头的“管理手册”下载。",
            "生产使用建议建立变更登记：操作人、时间、路径、修改前后摘要和回滚方式。",
        ],
    },
]


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_font(paragraph, size=11, color=INK, bold=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def add_page_number(paragraph):
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("Moonly Firebase Data Admin Guide")
    set_run_font(run, size=9, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(fp)


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    set_table_borders(table, color="E3E6EA")
    cell = table.cell(0, 0)
    set_cell_shading(cell, SOFT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_data_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "项目"
    headers[1].text = "说明"
    for cell in headers:
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        set_paragraph_font(p, size=9.5, color=DARK_BLUE, bold=True)
        p.paragraph_format.space_after = Pt(0)

    for label, detail in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
        set_paragraph_font(cells[0].paragraphs[0], size=9.5, color=INK, bold=True)
        set_paragraph_font(cells[1].paragraphs[0], size=9.5, color=INK)
        for cell in cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15

    set_table_geometry(table, [2400, 6960], indent_dxa=120)
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_bullets(doc, bullets):
    for item in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run_font(run, size=10.5, color=INK)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(54)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Moonly Firebase")
    set_run_font(run, size=14, color=MUTED, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("网页版数据管理系统手册")
    set_run_font(run, size=25, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("可获取、可修改的 Firebase 内容清单与后台操作规范")
    set_run_font(run, size=13, color=DARK_BLUE)

    add_data_table(doc, [
        ("文档版本", "v1.0"),
        ("适用项目", "fari-app-b2fd2 / MoonlyApp"),
        ("后台页面", "web-admin/data-admin.html；部署后 https://fari-app-b2fd2.web.app/data-admin"),
        ("后端接口", "adminDataCollections、adminDataList、adminDataUpsert、adminDataDelete"),
        ("最后更新", "2026-06-23"),
    ])

    add_callout(
        doc,
        "这份 20 页手册描述管理员可以通过网页版后台读取或修改的 Firestore 数据、字段边界、风险点和上线验收方式。"
    )

    doc.add_heading("阅读方式", level=2)
    add_bullets(doc, [
        "第 2-4 页说明系统能力、登录和集合总览。",
        "第 5-18 页逐项列出可管理的数据集合和可修改内容。",
        "第 19-20 页说明 JSON 特殊类型、部署和验收清单。",
    ])


def add_content_page(doc, page):
    doc.add_heading(page["title"], level=1)
    add_callout(doc, page["lead"])
    add_data_table(doc, page["table"])
    doc.add_heading("操作要点", level=2)
    add_bullets(doc, page["bullets"])


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_cover(doc)
    doc.add_page_break()
    for index, page in enumerate(PAGES):
        add_content_page(doc, page)
        if index != len(PAGES) - 1:
            doc.add_page_break()
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
